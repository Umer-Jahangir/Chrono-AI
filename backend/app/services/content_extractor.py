import csv
import io
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class UnsupportedContentType(ValueError):
    pass


def _normalize(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r" +", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text(data: bytes, filename: str, mime_type: str | None) -> str:
    """Extract searchable text from common Google Drive file formats."""
    mime = (mime_type or "").split(";", 1)[0].lower()
    suffix = Path(filename).suffix.lower()

    if mime == "application/pdf" or suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        return _normalize("\n\n".join(page.extract_text() or "" for page in reader.pages))

    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or suffix == ".docx":
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return _normalize("\n".join(parts))

    if mime == "application/json" or suffix == ".json":
        parsed = json.loads(data.decode("utf-8-sig"))
        return _normalize(json.dumps(parsed, ensure_ascii=False, indent=2))

    if mime in {"text/csv", "application/csv"} or suffix == ".csv":
        rows = csv.reader(io.StringIO(data.decode("utf-8-sig", errors="replace")))
        return _normalize("\n".join("\t".join(row) for row in rows))

    if mime.startswith("text/") or suffix in {".txt", ".md", ".log", ".xml", ".html", ".htm", ".yaml", ".yml"}:
        return _normalize(data.decode("utf-8-sig", errors="replace"))

    raise UnsupportedContentType(f"Unsupported file type: {mime or suffix or 'unknown'}")
